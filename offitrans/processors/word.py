"""
Word document processor for Offitrans

This module provides functionality to translate Word documents while preserving
formatting and layout.
"""

import logging
from typing import List, Dict, Any
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from .base import BaseProcessor
from ..exceptions.errors import WordProcessorError

logger = logging.getLogger(__name__)


class WordProcessor(BaseProcessor):
    """
    Word document processor that handles translation while preserving formatting.

    This processor can handle:
    - Paragraph text
    - Text formatting (bold, italic, etc.)
    - Headers and footers
    - Tables
    - Lists
    """

    def __init__(self, **kwargs):
        """
        Initialize Word processor.

        Args:
            **kwargs: Additional arguments passed to BaseProcessor
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise WordProcessorError(
                "python-docx library is required for Word processing",
                details="Install with: pip install python-docx",
            )

        super().__init__(**kwargs)

    def supports_file_type(self, file_path: str) -> bool:
        """
        Check if file type is supported.

        Args:
            file_path: Path to the file

        Returns:
            True if file type is supported
        """
        supported_extensions = {".docx", ".doc"}
        return Path(file_path).suffix.lower() in supported_extensions

    def extract_text(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text content from Word document.

        Args:
            file_path: Path to the Word document

        Returns:
            List of dictionaries containing text and metadata
        """
        text_data = []

        try:
            doc = Document(file_path)
            logger.info(f"Successfully opened Word document: {file_path}")

            # Extract text from paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Extract paragraph-level formatting
                    para_format = self._extract_paragraph_format(paragraph)

                    # Extract run-level formatting
                    runs_info = []
                    for run_idx, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            run_format = self._extract_run_format(run)
                            runs_info.append(
                                {
                                    "text": run.text,
                                    "run_index": run_idx,
                                    "format": run_format,
                                }
                            )

                    text_data.append(
                        {
                            "text": paragraph.text,
                            "type": "paragraph",
                            "paragraph_index": para_idx,
                            "paragraph_format": para_format,
                            "runs_info": runs_info,
                        }
                    )

                    logger.debug(
                        f"Extracted paragraph {para_idx}: '{paragraph.text[:50]}...'"
                    )

            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        if cell.text.strip():
                            text_data.append(
                                {
                                    "text": cell.text,
                                    "type": "table_cell",
                                    "table_index": table_idx,
                                    "row_index": row_idx,
                                    "cell_index": cell_idx,
                                }
                            )

                            logger.debug(
                                f"Extracted table cell [{table_idx}][{row_idx}][{cell_idx}]: '{cell.text[:50]}...'"
                            )

            logger.info(f"Total extracted {len(text_data)} text elements")
            return text_data

        except Exception as e:
            raise WordProcessorError(
                f"Failed to extract text from Word document",
                details=str(e),
                file_path=file_path,
            ) from e

    def translate_and_save(
        self, file_path: str, output_path: str, target_language: str = "en"
    ) -> bool:
        """
        Translate Word document and save to output path.

        Args:
            file_path: Path to input Word document
            output_path: Path for output Word document
            target_language: Target language code

        Returns:
            True if successful, False otherwise
        """
        try:
            # Step 1: Extract text and metadata
            logger.info("Step 1: Extracting text from Word document...")
            text_data = self.extract_text(file_path)

            if not text_data:
                logger.warning("No translatable text found in Word document")
                return False

            # Step 2: Preprocess and translate texts
            logger.info("Step 2: Translating texts...")
            original_texts = [item["text"] for item in text_data]
            unique_texts, metadata = self.preprocess_texts(original_texts)
            translated_unique = self.translate_texts(unique_texts, target_language)
            translated_texts = self.postprocess_translations(
                original_texts, translated_unique, metadata
            )

            # Step 3: Apply translations to Word document
            logger.info("Step 3: Applying translations to Word document...")
            success = self._replace_text_with_format(
                file_path, output_path, text_data, translated_texts, target_language
            )

            if success:
                logger.info(f"Successfully translated Word document: {output_path}")
                return True
            else:
                logger.error("Failed to apply translations to Word document")
                return False

        except Exception as e:
            logger.error(f"Error translating Word document: {e}")
            return False

    def _replace_text_with_format(
        self,
        word_path: str,
        output_path: str,
        text_data: List[Dict[str, Any]],
        translated_texts: List[str],
        target_language: str = "en",
    ) -> bool:
        """
        Replace text in Word document while preserving formatting.

        Args:
            word_path: Input Word document path
            output_path: Output Word document path
            text_data: Original text data with metadata
            translated_texts: List of translated texts
            target_language: Target language code

        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document(word_path)

            paragraph_translations = {}
            table_translations = {}

            # Group translations by type
            for item, translated_text in zip(text_data, translated_texts):
                if item["type"] == "paragraph":
                    paragraph_translations[item["paragraph_index"]] = {
                        "text": translated_text,
                        "format": item.get("paragraph_format", {}),
                        "runs_info": item.get("runs_info", []),
                    }
                elif item["type"] == "table_cell":
                    key = (item["table_index"], item["row_index"], item["cell_index"])
                    table_translations[key] = translated_text

            # Apply paragraph translations
            for para_idx, para_info in paragraph_translations.items():
                if para_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[para_idx]

                    # Clear existing text
                    paragraph.clear()

                    # Add translated text with formatting
                    run = paragraph.add_run(para_info["text"])
                    self._apply_run_format(
                        run, para_info.get("format", {}), target_language
                    )

                    logger.debug(f"Applied translation to paragraph {para_idx}")

            # Apply table cell translations
            for (
                table_idx,
                row_idx,
                cell_idx,
            ), translated_text in table_translations.items():
                if (
                    table_idx < len(doc.tables)
                    and row_idx < len(doc.tables[table_idx].rows)
                    and cell_idx < len(doc.tables[table_idx].rows[row_idx].cells)
                ):

                    cell = doc.tables[table_idx].rows[row_idx].cells[cell_idx]
                    cell.text = translated_text

                    logger.debug(
                        f"Applied translation to table cell [{table_idx}][{row_idx}][{cell_idx}]"
                    )

            # Save the document
            doc.save(output_path)

            logger.info(f"Successfully saved translated Word document: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Error replacing text in Word document: {e}")
            return False

    def _extract_paragraph_format(self, paragraph) -> Dict[str, Any]:
        """
        Extract formatting information from a paragraph.

        Args:
            paragraph: python-docx paragraph object

        Returns:
            Dictionary containing format information
        """
        format_info = {}

        try:
            if paragraph.style:
                format_info["style_name"] = paragraph.style.name

            # Alignment
            if paragraph.alignment:
                format_info["alignment"] = paragraph.alignment

            # Other paragraph formatting can be added here

        except Exception as e:
            logger.error(f"Error extracting paragraph format: {e}")

        return format_info

    def _extract_run_format(self, run) -> Dict[str, Any]:
        """
        Extract formatting information from a run.

        Args:
            run: python-docx run object

        Returns:
            Dictionary containing format information
        """
        format_info = {}

        try:
            if run.font:
                format_info["font_name"] = run.font.name
                format_info["font_size"] = run.font.size
                format_info["bold"] = run.font.bold
                format_info["italic"] = run.font.italic
                format_info["underline"] = run.font.underline
                format_info["color"] = run.font.color

        except Exception as e:
            logger.error(f"Error extracting run format: {e}")

        return format_info

    def _apply_run_format(
        self, run, format_info: Dict[str, Any], target_language: str = "en"
    ) -> None:
        """
        Apply formatting to a run.

        Args:
            run: python-docx run object
            format_info: Format information dictionary
            target_language: Target language code
        """
        try:
            if not format_info or not run.font:
                return

            # Font name (with language-specific adjustments)
            if target_language == "th" and format_info.get("font_name"):
                # Use Thai-compatible font
                run.font.name = "TH SarabunPSK"
            elif format_info.get("font_name"):
                run.font.name = format_info["font_name"]

            # Font size (with adjustment)
            if format_info.get("font_size"):
                original_size = format_info["font_size"]
                if original_size:
                    adjusted_size = max(
                        Pt(6), Pt(original_size.pt * self.font_size_adjustment)
                    )
                    run.font.size = adjusted_size

            # Other formatting
            if format_info.get("bold") is not None:
                run.font.bold = format_info["bold"]
            if format_info.get("italic") is not None:
                run.font.italic = format_info["italic"]
            if format_info.get("underline") is not None:
                run.font.underline = format_info["underline"]
            if format_info.get("color"):
                run.font.color = format_info["color"]

        except Exception as e:
            logger.error(f"Error applying run format: {e}")


# Function for simple Word document translation (backward compatibility)
def docx_translate(
    input_file: str, output_file: str, target_language: str = "en"
) -> bool:
    """
    Simple function to translate a Word document.

    Args:
        input_file: Path to input Word document
        output_file: Path to output Word document
        target_language: Target language code

    Returns:
        True if successful, False otherwise
    """
    try:
        processor = WordProcessor()
        return processor.process_file(input_file, output_file, target_language)
    except Exception as e:
        logger.error(f"Error in docx_translate: {e}")
        return False
