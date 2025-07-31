import os
import sys
from typing import List, Dict, Union

# 添加父目录到路径以导入翻译工具
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from translate_tools.google_translate import GoogleTranslator
from translate_tools.sungrow_translate import SunTranslator
from translate_tools.base import Translator



class DocxTranslator:
    """DOCX文档翻译器，支持保持样式的翻译，包括表格"""
    
    def __init__(self, source_lang: str = "zh", target_lang: str = "en", 
                 max_workers: int = 5, translator_type: str = "sungrow", **kwargs):
        """
        初始化DOCX翻译器
        
        Args:
            source_lang: 源语言代码
            target_lang: 目标语言代码
            max_workers: 并发工作者数量
            translator_type: 翻译器类型 ("google", "sungrow")
        """
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.max_workers = max_workers
        self.translator_type = translator_type
        
        # 根据类型初始化翻译器
        self.translator = self._init_translator(translator_type, **kwargs)
        
        # 缓存翻译结果
        self.translation_cache = {}
    
    def _init_translator(self, translator_type: str, **kwargs) -> Translator:
        """
        根据类型初始化翻译器
        
        Args:
            translator_type: 翻译器类型
            **kwargs: 额外参数
            
        Returns:
            翻译器实例
        """
        translator_classes = {
            "google": GoogleTranslator,
            "sungrow": SunTranslator,
        }
        
        if translator_type not in translator_classes:
            raise ValueError(f"不支持的翻译器类型: {translator_type}. 支持的类型: {list(translator_classes.keys())}")
        
        translator_class = translator_classes[translator_type]
        return translator_class(
            source_lang=self.source_lang,
            target_lang=self.target_lang,
            max_workers=self.max_workers,
            **kwargs
        )
        
    def _detect_language(self, text: str) -> str:
        """
        检测文本语言
        
        Args:
            text: 输入文本
            
        Returns:
            语言代码
        """
        if not text:
            return 'en'
            
        # 检测中文字符
        chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
        # 检测泰文字符
        thai_chars = sum(1 for c in text if '\u0e00' <= c <= '\u0e7f')
        
        if chinese_chars > 0:
            return 'zh'
        elif thai_chars > 0:
            return 'th'
        else:
            return 'en'
    
    def _should_translate(self, text: str) -> bool:
        """
        判断文本是否需要翻译
        
        Args:
            text: 文本内容
            
        Returns:
            是否需要翻译
        """
        if not text or not text.strip():
            return False
            
        # 如果文本只包含数字、标点符号或空白字符，不翻译
        if text.strip().replace(' ', '').replace('\n', '').replace('\t', '') == '':
            return False
            
        # 检查是否为纯数字
        if text.strip().replace('.', '').replace(',', '').replace('-', '').isdigit():
            return False
            
        # 检测语言，如果已经是目标语言则不翻译
        detected_lang = self._detect_language(text)
        if detected_lang == self.target_lang:
            return False
            
        return True
    
    def translate_text(self, text: str) -> str:
        """
        翻译单个文本，使用缓存
        
        Args:
            text: 要翻译的文本
            
        Returns:
            翻译后的文本
        """
        if not self._should_translate(text):
            return text
            
        # 检查缓存
        if text in self.translation_cache:
            return self.translation_cache[text]
            
        try:
            translated = self.translator.translate_text(text)
            self.translation_cache[text] = translated
            return translated
        except Exception as e:
            print(f"翻译失败: {text} - {str(e)}")
            return text
    
    def translate_batch(self, texts: List[str]) -> List[str]:
        """
        批量翻译文本
        
        Args:
            texts: 文本列表
            
        Returns:
            翻译结果列表
        """
        # 过滤需要翻译的文本
        texts_to_translate = []
        indices_to_translate = []
        
        for i, text in enumerate(texts):
            if self._should_translate(text):
                if text not in self.translation_cache:
                    texts_to_translate.append(text)
                    indices_to_translate.append(i)
        
        # 批量翻译
        if texts_to_translate:
            try:
                translated_texts = self.translator.translate_text_batch(texts_to_translate)
                # 更新缓存
                for original, translated in zip(texts_to_translate, translated_texts):
                    self.translation_cache[original] = translated
            except Exception as e:
                print(f"批量翻译失败: {str(e)}")
        
        # 构建结果列表
        results = []
        for text in texts:
            if self._should_translate(text):
                results.append(self.translation_cache.get(text, text))
            else:
                results.append(text)
                
        return results
    
    def _extract_run_properties(self, run: Run) -> Dict:
        """
        提取Run的格式属性
        
        Args:
            run: docx Run对象
            
        Returns:
            格式属性字典
        """
        properties = {}
        
        if run.bold is not None:
            properties['bold'] = run.bold
        if run.italic is not None:
            properties['italic'] = run.italic
        if run.underline is not None:
            properties['underline'] = run.underline
        if run.font.size is not None:
            properties['size'] = run.font.size
        if run.font.name is not None:
            properties['name'] = run.font.name
        if run.font.color.rgb is not None:
            properties['color'] = run.font.color.rgb
            
        return properties
    
    def _apply_run_properties(self, run: Run, properties: Dict):
        """
        应用格式属性到Run
        
        Args:
            run: docx Run对象
            properties: 格式属性字典
        """
        if 'bold' in properties:
            run.bold = properties['bold']
        if 'italic' in properties:
            run.italic = properties['italic']
        if 'underline' in properties:
            run.underline = properties['underline']
        if 'size' in properties:
            run.font.size = properties['size']
        if 'name' in properties:
            run.font.name = properties['name']
        if 'color' in properties:
            run.font.color.rgb = properties['color']
    
    def _translate_paragraph(self, paragraph: Paragraph):
        """
        翻译段落，保持格式
        
        Args:
            paragraph: docx段落对象
        """
        if not paragraph.text.strip():
            return
            
        # 收集所有run的文本和属性
        runs_info = []
        for run in paragraph.runs:
            if run.text:
                runs_info.append({
                    'text': run.text,
                    'properties': self._extract_run_properties(run)
                })
        
        if not runs_info:
            return
            
        # 提取所有文本用于翻译
        texts_to_translate = [info['text'] for info in runs_info]
        
        # 批量翻译
        translated_texts = self.translate_batch(texts_to_translate)
        
        # 清除段落中的所有run
        for run in paragraph.runs[:]:
            run._element.getparent().remove(run._element)
        
        # 重新创建run并应用格式
        for i, (translated_text, run_info) in enumerate(zip(translated_texts, runs_info)):
            if translated_text:
                new_run = paragraph.add_run(translated_text)
                self._apply_run_properties(new_run, run_info['properties'])
    
    def _translate_table(self, table: Table):
        """
        翻译表格，保持样式
        
        Args:
            table: docx表格对象
        """
        print(f"正在翻译表格，行数: {len(table.rows)}, 列数: {len(table.columns) if table.rows else 0}")
        
        # 收集所有单元格文本
        all_texts = []
        cell_info = []
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        # 收集段落中所有run的信息
                        runs_info = []
                        for run in paragraph.runs:
                            if run.text:
                                runs_info.append({
                                    'text': run.text,
                                    'properties': self._extract_run_properties(run)
                                })
                                all_texts.append(run.text)
                        
                        if runs_info:
                            cell_info.append({
                                'row': row_idx,
                                'col': col_idx,
                                'paragraph': paragraph,
                                'runs_info': runs_info
                            })
        
        if not all_texts:
            return
            
        print(f"找到 {len(all_texts)} 个文本片段需要翻译")
        
        # 批量翻译所有文本
        translated_texts = self.translate_batch(all_texts)
        
        # 应用翻译结果
        text_idx = 0
        for cell_data in cell_info:
            paragraph = cell_data['paragraph']
            runs_info = cell_data['runs_info']
            
            # 清除段落中的所有run
            for run in paragraph.runs[:]:
                run._element.getparent().remove(run._element)
            
            # 重新创建run并应用翻译和格式
            for run_info in runs_info:
                translated_text = translated_texts[text_idx]
                if translated_text:
                    new_run = paragraph.add_run(translated_text)
                    self._apply_run_properties(new_run, run_info['properties'])
                text_idx += 1
        
        print(f"表格翻译完成")
    
    def translate_docx(self, input_path: str, output_path: str):
        """
        翻译DOCX文档
        
        Args:
            input_path: 输入文档路径
            output_path: 输出文档路径
        """
        print(f"开始翻译文档: {input_path}")
        
        try:
            # 打开文档
            doc = Document(input_path)
            
            # 统计信息
            paragraph_count = 0
            table_count = 0
            
            # 翻译所有段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    self._translate_paragraph(paragraph)
                    paragraph_count += 1
            
            # 翻译所有表格
            for table in doc.tables:
                self._translate_table(table)
                table_count += 1
            
            # 保存文档
            doc.save(output_path)
            
            print(f"翻译完成!")
            print(f"- 处理段落: {paragraph_count}")
            print(f"- 处理表格: {table_count}")
            print(f"- 输出文件: {output_path}")
            
        except Exception as e:
            print(f"翻译文档时出错: {str(e)}")
            raise
    
    def batch_translate_docx(self, input_folder: str, output_folder: str, 
                           file_pattern: str = "*.docx"):
        """
        批量翻译DOCX文档
        
        Args:
            input_folder: 输入文件夹路径
            output_folder: 输出文件夹路径
            file_pattern: 文件匹配模式
        """
        import glob
        
        # 确保输出文件夹存在
        os.makedirs(output_folder, exist_ok=True)
        
        # 查找所有匹配的文件
        pattern = os.path.join(input_folder, file_pattern)
        docx_files = glob.glob(pattern)
        
        if not docx_files:
            print(f"在 {input_folder} 中未找到匹配 {file_pattern} 的文件")
            return
        
        print(f"找到 {len(docx_files)} 个文档需要翻译")
        
        for i, docx_file in enumerate(docx_files, 1):
            try:
                filename = os.path.basename(docx_file)
                name, ext = os.path.splitext(filename)
                output_filename = f"{name}_translated_{self.target_lang}{ext}"
                output_path = os.path.join(output_folder, output_filename)
                
                print(f"\n[{i}/{len(docx_files)}] 正在处理: {filename}")
                
                self.translate_docx(docx_file, output_path)
                
                print(f"✅ 完成: {output_filename}")
                
            except Exception as e:
                print(f"❌ 处理 {docx_file} 时出错: {str(e)}")
                continue
        
        print(f"\n批量翻译完成！输出目录: {output_folder}")
    
    def get_translation_stats(self) -> Dict:
        """
        获取翻译统计信息
        
        Returns:
            统计信息字典
        """
        return {
            'cache_size': len(self.translation_cache),
            'source_lang': self.source_lang,
            'target_lang': self.target_lang,
            'max_workers': self.max_workers,
            'translator_type': self.translator_type
        }
    
    def switch_translator(self, translator_type: str, **kwargs):
        """
        切换翻译器类型
        
        Args:
            translator_type: 翻译器类型
            **kwargs: 额外参数
        """
        self.translator_type = translator_type
        self.translator = self._init_translator(translator_type, **kwargs)
        # 清空缓存，因为翻译器变了
        self.translation_cache.clear()
        print(f"已切换到 {translator_type} 翻译器")
    
    @staticmethod
    def get_available_translators() -> List[str]:
        """
        获取可用的翻译器类型
        
        Returns:
            可用翻译器列表
        """
        return ["google", "sungrow"]
    
    def test_translator(self) -> bool:
        """
        测试当前翻译器是否可用
        
        Returns:
            是否可用
        """
        try:
            test_text = "测试"
            result = self.translator.translate_text(test_text)
            return result != test_text
        except Exception as e:
            print(f"翻译器测试失败: {str(e)}")
            return False


def create_docx_translator(translator_type: str = "google", 
                          source_lang: str = "zh", 
                          target_lang: str = "en",
                          max_workers: int = 5,
                          **kwargs) -> DocxTranslator:
    """
    工厂方法：创建DOCX翻译器
    
    Args:
        translator_type: 翻译器类型 ("google", "sungrow")
        source_lang: 源语言代码
        target_lang: 目标语言代码
        max_workers: 并发工作者数量
        **kwargs: 额外参数
        
    Returns:
        DocxTranslator实例
    """
    return DocxTranslator(
        source_lang=source_lang,
        target_lang=target_lang,
        max_workers=max_workers,
        translator_type=translator_type,
        **kwargs
    )


def translate_single_docx(input_path: str, 
                         output_path: str,
                         translator_type: str = "google",
                         source_lang: str = "zh", 
                         target_lang: str = "en",
                         **kwargs):
    """
    便捷函数：翻译单个DOCX文档
    
    Args:
        input_path: 输入文档路径
        output_path: 输出文档路径
        translator_type: 翻译器类型
        source_lang: 源语言代码
        target_lang: 目标语言代码
        **kwargs: 额外参数
    """
    translator = create_docx_translator(
        translator_type=translator_type,
        source_lang=source_lang,
        target_lang=target_lang,
        **kwargs
    )
    
    # 测试翻译器
    if not translator.test_translator():
        print(f"❌ {translator_type} 翻译器不可用")
        return False
        
    translator.translate_docx(input_path, output_path)
    return True


def main():
    """主函数 - 使用示例"""
    print("=== DOCX翻译器使用示例 ===")
    
    # 显示可用的翻译器
    available_translators = DocxTranslator.get_available_translators()
    print(f"可用翻译器: {available_translators}")
    
    # 初始化翻译器（默认使用Google翻译）
    translator = DocxTranslator(
        source_lang="zh",     # 中文
        target_lang="en",     # 英文
        max_workers=5,
        translator_type="sungrow"  # 默认使用Sungrow翻译
    )
    
    # 测试翻译器
    # print(f"\n测试Google翻译器...")
    # if translator.test_translator():
    #     print("✅ Google翻译器可用")
    # else:
    #     print("❌ Google翻译器不可用，切换到Sungrow翻译器")
    #     translator.switch_translator("sungrow")
    #     if translator.test_translator():
    #         print("✅ Sungrow翻译器可用")
    #     else:
    #         print("❌ 所有翻译器都不可用")
    #         return
    
    # 单文件翻译示例
    # translator.translate_docx(
    #     input_path="A0SK0060-datasheet (2).docx",
    #     output_path="A0SK0060-datasheet (2)_translated_en.docx"
    # )
    
    # 批量翻译示例
    translator.batch_translate_docx(
        input_folder="input_docs",
        output_folder="output_docs",
        file_pattern="*.docx"
    )
    
    # 演示切换翻译器
    # print("\n切换到Sungrow翻译器...")
    # translator.switch_translator("sungrow")
    # translator.translate_docx(
    #     input_path="test2.docx",
    #     output_path="test2_translated_sungrow.docx"
    # )
    
    # 打印统计信息
    stats = translator.get_translation_stats()
    print("\n翻译器配置:")
    for key, value in stats.items():
        print(f"  {key}: {value}")
    
    print("\n=== 使用示例完成 ===")


if __name__ == "__main__":
    main()